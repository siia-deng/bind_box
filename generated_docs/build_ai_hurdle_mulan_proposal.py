from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "/Users/alchemistthe/Documents/hurdle-club/website/generated_docs/AI跨栏 x Mulan AI视频创作平台活动提案V1.docx"


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_para_format(paragraph, before=0, after=8, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="DADCE0", size="4"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for p in cell.paragraphs:
                set_para_format(p, after=0, line=1.15)
                for run in p.runs:
                    set_run_font(run, size=10)


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 3"]
    set_para_format(p, before=16, after=4, line=1.15)
    r = p.add_run(text)
    set_run_font(r, size=14, color="434343")
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    set_para_format(p)
    r = p.add_run(text)
    set_run_font(r, size=11)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_para_format(p, after=4, line=1.15)
    r = p.add_run(text)
    set_run_font(r, size=11)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    set_para_format(p, after=4, line=1.15)
    r = p.add_run(text)
    set_run_font(r, size=11)
    return p


def style_table_header(row):
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                set_run_font(run, size=10, bold=True)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    for name in ("Normal", "List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(11)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_after = Pt(8 if name == "Normal" else 4)
        style.paragraph_format.line_spacing = 1.15

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_format(title, after=3, line=1.15)
    r = title.add_run("AI跨栏 x Mulan AI视频创作平台线下 AI 视觉 Workshop 活动提案 (V1)")
    set_run_font(r, size=26)

    add_heading(doc, "1. 合作背景与愿景")
    add_bullet(doc, "AI跨栏社群：面向 AI 创作者、品牌主理人、内容从业者与视觉表达爱好者，由社群创始人、AI 视觉制作专业者亲授线下 AI 视觉 workshop 教学方案，把方法论学习转化为可发布的视频作品。")
    add_bullet(doc, "Mulan AI视频创作平台：提供 AI 视频生成与创作能力，并以算力积分作为用户使用和体验平台的核心权益载体。")
    add_bullet(doc, "核心目标：通过“创始人亲授方法论 + Mulan 平台实操 + 平台积分 + 作品发布 + 数据归因”的闭环，让参与者用 AI跨栏的方法在 Mulan 上搭建创作工作流并产出 AI 视频，同时为 Mulan 带来可追踪的新用户、内容曝光与后续付费转化数据。")

    add_heading(doc, "2. 活动主题：用 Mulan 把灵感变成可发布的 AI 视频")
    add_bullet(doc, "备选主题：")
    add_bullet(doc, "“一场 workshop，完成你的第一支 AI 视频作品”。")
    add_bullet(doc, "“从灵感到发布：AI 视觉创作实战工作坊”。")
    add_bullet(doc, "“让创作者现场用 Mulan 做出可传播内容”。")

    add_heading(doc, "3. 核心机制设计")
    add_body(doc, "为了形成线下活动收入、平台真实使用、社交传播和后续付费转化的增长闭环，建议设计以下机制：")
    add_bullet(doc, "积分采购机制：AI跨栏社群根据每场线下 AI 视觉 workshop 的实际参与人数，向 Mulan 采购活动使用的算力积分；当前建议标准为 500 积分/人，该积分费用包含在 AI跨栏活动门票中。")
    add_bullet(doc, "专属邀请码与用户归因：Mulan 为 AI跨栏社群提供专属邀请码或专属注册链接。参与者通过该邀请码领取活动积分后，Mulan 可在后台识别其来源为 AI跨栏社群，用于统计新增用户、激活、使用和后续付费转化。")
    add_bullet(doc, "现场转化路径：活动现场由 AI跨栏社群创始人亲授 AI 视觉制作方法、案例拆解和工作流搭建方案；参与者注册 Mulan、绑定专属邀请码、领取 500 积分，并在 Mulan 平台上按照该方法完成至少一支 AI 视频作品。")
    add_bullet(doc, "社交传播激励：活动结束后，AI跨栏邀请参与者将视频作品发布至小红书等社交媒体，并添加话题 #AI跨栏 #mulan；作品需保留一周，并将发布链接回传给 AI跨栏进行记录。")
    add_bullet(doc, "二次积分奖励：对完成发布、话题标记、链接回传且保留满一周的参与者，Mulan 可在其账号中追加发放 100 积分，作为社交传播与复访创作激励。")
    add_bullet(doc, "数据复盘闭环：每场活动结束后，双方可复盘报名人数、邀请码领取人数、作品完成人数、社交发布人数、曝光与互动、追加积分领取人数，以及后续在 Mulan 的付费转化率。")

    add_heading(doc, "4. 活动流程与日程计划")
    table = doc.add_table(rows=1, cols=3)
    headers = ["阶段", "时间", "关键任务"]
    for i, text in enumerate(headers):
        table.cell(0, i).text = text
    style_table_header(table.rows[0])
    rows = [
        ("合作确认期", "活动前 2-3 周", "确认积分采购单价与 500 积分/人的发放标准；确认专属 AI跨栏社群邀请码、后台归因口径、结算方式与数据看板字段。"),
        ("报名与宣发期", "活动前 1-2 周", "AI跨栏负责活动招募、社群预热、门票售卖与参与者通知；Mulan 可提供平台介绍、案例素材、邀请码说明与联合宣发文案。"),
        ("现场执行期", "活动当天", "参与者使用专属邀请码注册并领取积分；AI跨栏创始人亲授 AI 视觉制作方法、工作流搭建与作品辅导，带参与者在 Mulan 平台产出 AI 视频；Mulan 支持账号、积分、生成额度等平台问题。"),
        ("发布裂变期", "活动后 1 周", "AI跨栏邀请参与者发布作品至小红书等社交媒体，添加 #AI跨栏 #mulan，并回传链接；双方记录内容样本与传播数据。"),
        ("奖励与复盘期", "活动后第 7-10 天", "核验作品保留满一周后，Mulan 向符合条件的用户追加 100 积分；双方复盘新增、活跃、UGC、转化率与下一场活动优化。"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
    set_table_width(table, [1.15, 1.35, 4.0])

    add_heading(doc, "5. 关键考核指标 (KPI)")
    add_bullet(doc, "活动报名与到场：每场报名人数、到场人数、实际积分发放人数。")
    add_bullet(doc, "平台新增与激活：通过 AI跨栏专属邀请码注册的用户数、领取 500 积分用户数、首次生成用户数。")
    add_bullet(doc, "作品完成：现场完成 AI 视频作品的人数、作品数量、平均消耗积分。")
    add_bullet(doc, "社交传播：发布至小红书等社交媒体的作品数量、话题 #AI跨栏 #mulan 覆盖量、互动数、回传链接数。")
    add_bullet(doc, "复访与留存：完成发布后领取 100 积分的人数、活动后 7 日/30 日再次使用 Mulan 的用户比例。")
    add_bullet(doc, "商业转化：来自 AI跨栏渠道用户在 Mulan 的后续付费人数、付费金额、付费转化率与客单价。")

    add_heading(doc, "6. 双方分工")
    table2 = doc.add_table(rows=1, cols=3)
    for i, text in enumerate(["事项", "AI跨栏社群", "Mulan"]):
        table2.cell(0, i).text = text
    style_table_header(table2.rows[0])
    rows2 = [
        ("活动组织", "由 AI跨栏社群创始人负责 workshop 策划、招募、门票售卖、现场亲授、作品辅导与社群运营。", "提供平台能力说明、账号/积分支持与必要的产品资料。"),
        ("积分与结算", "按实际参与人数采购 500 积分/人，并将积分成本纳入活动门票。", "按约定向专属邀请码用户发放活动积分，并支持追加 100 积分奖励。"),
        ("用户归因", "引导参与者使用专属 AI跨栏社群邀请码注册或兑换。", "提供渠道归因与用户数据，支持查看来自 AI跨栏渠道的后续使用和付费情况。"),
        ("传播裂变", "组织参与者发布作品、添加话题、回传链接并完成一周保留提醒。", "配合奖励发放，可选择转发优质作品或沉淀联合案例。"),
    ]
    for row in rows2:
        cells = table2.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
    set_table_width(table2, [1.2, 2.6, 2.7])

    add_heading(doc, "7. 核心优势总结")
    add_bullet(doc, "专业方法论带动平台使用：活动不是单纯工具演示，而是由 AI跨栏创始人以成熟 AI 视觉制作方法带领参与者在 Mulan 上搭建工作流并完成作品，更容易形成深度使用。")
    add_bullet(doc, "真实使用场景：参与者不是单纯领取福利，而是在现场完成账号注册、积分使用、工作流搭建和作品产出，能更真实地验证平台体验。")
    add_bullet(doc, "可追踪增长链路：专属邀请码让 Mulan 可以清晰识别 AI跨栏渠道用户，并评估从活动参与到后续付费的完整转化率。")
    add_bullet(doc, "内容即传播：每位参与者产出的视频作品都可以成为平台案例与社交素材，#AI跨栏 #mulan 话题机制有助于形成持续曝光。")
    add_bullet(doc, "激励成本可控：基础 500 积分随门票采购，追加 100 积分仅面向完成发布并保留一周的用户发放，奖励与传播行为直接绑定。")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
